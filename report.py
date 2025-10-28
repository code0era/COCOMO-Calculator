# report.py
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import math
import os

def generate_report(project_name, project_type, kloc, results, cost_drivers, cocomo, graph_path=None):
    """
    Generate a professional COCOMO report in DOCX format with tables and optional graph.
    """

    doc = Document()

    # Set default font to Arial (san-serif)
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Section 1: Header
    header = doc.add_heading(f"COCOMO Report: {project_name}", level=0)
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_paragraph = doc.add_paragraph(f"Date & Time: {datetime.now().strftime('%d-%m-%Y %H:%M:%S')}")
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_paragraph("\n")  # spacing

    # Section 2: Project Info
    doc.add_heading("PROJECT INFORMATION", level=1)
    doc.add_paragraph(f"Project Name: {project_name}")
    doc.add_paragraph(f"Type of Project: {project_type}")
    doc.add_paragraph(f"KLOC: {kloc}")

    # ABCD coefficients
    A, B, C, D = cocomo.coefficients[project_type]
    doc.add_paragraph("Coefficients Used:")
    doc.add_paragraph(f"A = {A}, B = {B}, C = {C}, D = {D}")

    doc.add_paragraph("\n")  # spacing

    # Section 3: Outputs Table
    doc.add_heading("OUTPUTS", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'  # Safe built-in style

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Calculated'
    hdr_cells[2].text = 'Ceiling Value'

    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            run = paragraph.runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # black text

    for key in ['Effort (PM)', 'Development Time (Months)', 'Average Staff']:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(round(results[key], 2))
        row_cells[2].text = str(math.ceil(results[key]))

    # Formulas
    doc.add_paragraph("\nFormulas Used:")
    doc.add_paragraph(f"Effort = A * (KLOC ^ B) * EAF = {A} * ({kloc} ^ {B}) * EAF")
    doc.add_paragraph(f"Duration = C * (Effort ^ D) = {C} * (Effort ^ {D})")
    doc.add_paragraph("Staff = Effort / Duration")

    # Optional Summary
    doc.add_paragraph(f"\nSummary: The project requires approximately {math.ceil(results['Effort (PM)'])} "
                      f"person-months, with an estimated duration of {math.ceil(results['Development Time (Months)'])} months "
                      f"and average staffing of {math.ceil(results['Average Staff'])} persons.")

    doc.add_paragraph("\n")  # spacing

    # Section 4: Cost Drivers Table
    doc.add_heading("COST DRIVERS (EAF)", level=1)
    eaf_table = doc.add_table(rows=1, cols=2)
    eaf_table.style = 'Table Grid'  # Safe built-in style

    hdr_cells = eaf_table.rows[0].cells
    hdr_cells[0].text = 'Cost Driver'
    hdr_cells[1].text = 'Multiplier Value'

    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            run = paragraph.runs[0]
            run.font.name = 'Arial'
            run.font.bold = True

    for driver, val in cost_drivers.items():
        row_cells = eaf_table.add_row().cells
        row_cells[0].text = driver
        row_cells[1].text = str(val)

    # Section 5: Graph (optional)
    if graph_path and os.path.exists(graph_path):
        doc.add_paragraph("\n")
        doc.add_heading("Graphical Representation", level=1)
        doc.add_picture(graph_path, width=Inches(5))

    # Save document
    filename = f"COCOMO_Report_{project_name.replace(' ', '_')}.docx"
    doc.save(filename)
    return filename
